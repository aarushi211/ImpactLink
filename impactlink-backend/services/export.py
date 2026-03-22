"""
services/export.py

Service for exporting proposal sections to DOCX and PDF.
"""

import os
from docx import Document
from fpdf import FPDF
from typing import Dict, List

def export_to_docx(proposal_data: dict, output_path: str):
    """
    Exports proposal sections to a DOCX file.
    """
    doc = Document()
    
    # Title
    doc.add_heading(proposal_data.get("grant_title", "Grant Proposal"), 0)
    doc.add_paragraph(f"Organization: {proposal_data.get('org_name', 'N/A')}")
    doc.add_paragraph(f"Agency: {proposal_data.get('agency', 'N/A')}")
    doc.add_page_break()
    
    # Sections
    sections = proposal_data.get("sections", {})
    order = proposal_data.get("section_order", [])
    
    for key in order:
        if key in sections:
            section = sections[key]
            doc.add_heading(section.get("title", "Untitled Section"), level=1)
            doc.add_paragraph(section.get("content", ""))
            doc.add_page_break()
            
    doc.save(output_path)
    return output_path

def export_to_pdf(proposal_data: dict, output_path: str):
    """
    Exports proposal sections to a PDF file.
    """
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Add a Unicode font if possible, otherwise use standard fonts
    # fpdf2 supports standard fonts by default.
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    
    # Title
    pdf.cell(200, 10, txt=proposal_data.get("grant_title", "Grant Proposal"), ln=True, align='C')
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"Organization: {proposal_data.get('org_name', 'N/A')}", ln=True)
    pdf.cell(200, 10, txt=f"Agency: {proposal_data.get('agency', 'N/A')}", ln=True)
    pdf.ln(10)
    
    # Sections
    sections = proposal_data.get("sections", {})
    order = proposal_data.get("section_order", [])
    
    for key in order:
        if key in sections:
            section = sections[key]
            pdf.add_page()
            pdf.set_font("Arial", 'B', 14)
            pdf.cell(200, 10, txt=section.get("title", "Untitled Section"), ln=True)
            pdf.ln(5)
            pdf.set_font("Arial", size=11)
            # multi_cell for long text
            pdf.multi_cell(0, 10, txt=section.get("content", ""))
            
    pdf.output(output_path)
    return output_path
